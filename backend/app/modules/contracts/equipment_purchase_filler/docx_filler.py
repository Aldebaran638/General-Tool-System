"""
Equipment Purchase Contract Filler DOCX Logic

Hardcoded filling rules for the equipment purchase contract template.
"""

import json
import re
import tempfile
import uuid
from pathlib import Path

from docx import Document

from .schemas import ContractField, PreviewSegment


# =============================================================================
# Hardcoded field definitions for the equipment purchase contract template
# =============================================================================

CONTRACT_FIELDS: list[ContractField] = [
    ContractField(key="contract_number", label="合同编号"),
    ContractField(key="party_b_name", label="卖方名称"),
    ContractField(key="party_b_legal_person", label="卖方法定代表人"),
    ContractField(key="party_b_address", label="卖方地址"),
    ContractField(key="contract_amount_upper", label="合同总价大写"),
    ContractField(key="contract_amount_lower", label="合同总价小写"),
    ContractField(key="tax_rate", label="税率（%）"),
    ContractField(key="delivery_days_after_signing", label="合同签订后天数"),
    ContractField(key="delivery_days_after_prepayment", label="收到预付款后天数"),
    ContractField(key="delivery_location", label="交货地点"),
    ContractField(key="prepayment_days", label="预付款支付天数"),
    ContractField(key="prepayment_ratio", label="预付款比例（%）"),
    ContractField(key="prepayment_amount_upper", label="预付款大写金额"),
    ContractField(key="prepayment_amount_lower", label="预付款小写金额"),
    ContractField(key="acceptance_days", label="验收款支付天数"),
    ContractField(key="acceptance_ratio", label="验收款比例（%）"),
    ContractField(key="acceptance_amount_upper", label="验收款大写金额"),
    ContractField(key="acceptance_amount_lower", label="验收款小写金额"),
    ContractField(key="warranty_amount_upper", label="质保金大写金额"),
    ContractField(key="warranty_amount_lower", label="质保金小写金额"),
    ContractField(key="warranty_payment_days", label="质保期满后付款天数"),
    ContractField(key="invoice_type", label="票据类型"),
    ContractField(key="warranty_months", label="质量保证期月数"),
    ContractField(key="party_a_authorized_representative", label="买方授权代表"),
    ContractField(key="party_a_date", label="买方日期", type="date"),
    ContractField(key="party_b_authorized_representative", label="卖方授权代表"),
    ContractField(key="party_b_date", label="卖方日期", type="date"),
]

FIELD_KEYS = {f.key for f in CONTRACT_FIELDS}

EQUIPMENT_TABLE_COLUMNS = [
    "序号",
    "设备名称",
    "品牌",
    "规格",
    "单位",
    "数量",
    "单价（RMB）",
    "总价（RMB）",
]


# =============================================================================
# Preview extraction
# =============================================================================

_PREVIEW_MARKERS: list[tuple[str, str]] = [
    (r"^\s*(合同编号：)\s*$", r"\1{{contract_number}}"),
    (r"^(卖方：)\s*$", r"\1{{party_b_name}}"),
    (r"^(法定代表人：)\s*$", r"\1{{party_b_legal_person}}"),
    (r"^(地址：)\s*$", r"\1{{party_b_address}}"),
    (r"(合同总价为人民币：)\s*元整（￥\s*）", r"\1{{contract_amount_upper}}（￥{{contract_amount_lower}}）"),
    (r"(合同总价为含税价（税率：)\s*%", r"\1{{tax_rate}}%"),
    (r"(本合同签订后)\s*天之内", r"\1{{delivery_days_after_signing}}天之内"),
    (r"(收到预付款后)\s*天之内", r"\1{{delivery_days_after_prepayment}}天之内"),
    (r"(卖方将合同设备送到买方指定地点即)\s*。", r"\1{{delivery_location}}。"),
    (r"(合同签订后)\s*天内，买方支付合同总价\s*%即人民币\s*元整（￥\s*）的预付款",
     r"\1{{prepayment_days}}天内，买方支付合同总价{{prepayment_ratio}}%即人民币{{prepayment_amount_upper}}（￥{{prepayment_amount_lower}}）的预付款"),
    (r"(设备交付买方并安装调试，达到正常使用要求，经买方终验收合格后)\s*天内，买方支付合同总价\s*%即人民币\s*元整（￥\s*）。",
     r"\1{{acceptance_days}}天内，买方支付合同总价{{acceptance_ratio}}%即人民币{{acceptance_amount_upper}}（￥{{acceptance_amount_lower}}）。"),
    (r"(剩余合同总价5%即人民币)\s*元整（￥\s*）作为质保金",
     r"\1{{warranty_amount_upper}}（￥{{warranty_amount_lower}}）作为质保金"),
    (r"(质保期满后)\s*日之内扣除相关费用",
     r"\1{{warranty_payment_days}}日之内扣除相关费用"),
    (r"(卖方收取合同价款之前应向买方提供)\s*合法等额完税发票\s*/\s*合法等额收据",
     r"\1{{invoice_type}}"),
    (r"(设备的质量保证期为设备整机调试、试运行验收合格之日起)\s*个月",
     r"\1{{warranty_months}}个月"),
]


def _insert_preview_markers(text: str) -> str:
    """Replace blank placeholders in contract text with {{field_key}} markers."""
    for pattern, replacement in _PREVIEW_MARKERS:
        text = re.sub(pattern, replacement, text, count=1)
    return text


def extract_preview_with_markers(template_path: Path) -> list[PreviewSegment]:
    """Extract preview segments with {{field_key}} markers at blank positions."""
    doc = Document(str(template_path))
    segments: list[PreviewSegment] = []

    in_signing_section = False
    current_signing_party: str | None = None
    party_a_date_marker_done = False

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            continue

        stripped = text.strip()

        # Detect signing section start
        if "（本页无正文，仅为签署）" in text:
            in_signing_section = True
            current_signing_party = "A"
            segments.append(
                PreviewSegment(
                    type="paragraph",
                    text=text,
                    style=para.style.name if para.style else "Normal",
                )
            )
            continue

        if in_signing_section:
            if stripped.startswith("卖方："):
                current_signing_party = "B"
                text = re.sub(r"^(卖方：)\s*", r"\1{{party_b_name}}", text, count=1)
            elif stripped.startswith("授权代表："):
                if current_signing_party == "A":
                    text = re.sub(r"^(授权代表：)\s*", r"\1{{party_a_authorized_representative}}", text, count=1)
                else:
                    text = re.sub(r"^(授权代表：)\s*", r"\1{{party_b_authorized_representative}}", text, count=1)
            elif stripped == "日期：":
                if not party_a_date_marker_done:
                    text = re.sub(r"^(日期：)\s*", r"\1{{party_a_date}}", text, count=1)
                    party_a_date_marker_done = True
                else:
                    text = re.sub(r"^(日期：)\s*", r"\1{{party_b_date}}", text, count=1)
            elif stripped == "法定代表人：" and "申振威" not in text:
                text = re.sub(r"^(法定代表人：)\s*", r"\1{{party_b_legal_person}}", text, count=1)
            else:
                text = _insert_preview_markers(text)
        else:
            text = _insert_preview_markers(text)

        segments.append(
            PreviewSegment(
                type="paragraph",
                text=text,
                style=para.style.name if para.style else "Normal",
            )
        )

    # Tables: return as-is; the frontend will render editable cells for the equipment table
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        if rows:
            segments.append(PreviewSegment(type="table", rows=rows))

    return segments


def extract_preview(template_path: Path) -> list[PreviewSegment]:
    """Extract structured preview segments from the DOCX template."""
    return extract_preview_with_markers(template_path)


# =============================================================================
# Filling helpers
# =============================================================================

def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace all text in a paragraph while keeping the first run."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def _fill_equipment_table(doc, equipment_items: list[dict[str, str]]) -> None:
    """Fill the equipment detail table with the provided rows."""
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        if header_cells != EQUIPMENT_TABLE_COLUMNS:
            continue

        # Clear existing data rows (keep header)
        for row in reversed(table.rows[1:]):
            table._element.remove(row._element)

        # Add equipment rows
        for idx, item in enumerate(equipment_items, start=1):
            row = table.add_row()
            cells = row.cells
            cells[0].text = str(idx)
            cells[1].text = item.get("name", "")
            cells[2].text = item.get("brand", "")
            cells[3].text = item.get("spec", "")
            cells[4].text = item.get("unit", "")
            cells[5].text = item.get("quantity", "")
            cells[6].text = item.get("unit_price", "")
            cells[7].text = item.get("total_price", "")

        # If no items, add one empty row to keep table structure
        if not equipment_items:
            row = table.add_row()
            for cell in row.cells:
                cell.text = ""
        return


# =============================================================================
# Main fill function
# =============================================================================

def fill_contract(
    template_path: Path,
    field_values: dict[str, str],
    output_path: Path,
    equipment_items: list[dict[str, str]] | None = None,
) -> None:
    """Fill the equipment purchase contract template and save to output_path."""
    doc = Document(str(template_path))
    values = {k: v for k, v in field_values.items() if k in FIELD_KEYS}

    in_signing_section = False
    current_signing_party: str | None = None
    party_a_date_filled = False

    for p in doc.paragraphs:
        t = p.text
        stripped = t.strip()

        # Detect signing section start
        if "（本页无正文，仅为签署）" in t:
            in_signing_section = True
            current_signing_party = "A"
            continue

        if in_signing_section and stripped == "卖方：":
            current_signing_party = "B"
            _replace_paragraph_text(p, f"卖方：{values.get('party_b_name', '')}")
            continue

        # Contract number
        if not in_signing_section and stripped.startswith("合同编号："):
            _replace_paragraph_text(p, f"合同编号：{values.get('contract_number', '')}")
            continue

        # Party B info (body section)
        if not in_signing_section and stripped == "卖方：":
            _replace_paragraph_text(p, f"卖方：{values.get('party_b_name', '')}")
            continue

        if not in_signing_section and stripped == "法定代表人：" and "申振威" not in t:
            _replace_paragraph_text(p, f"法定代表人：{values.get('party_b_legal_person', '')}")
            continue

        if not in_signing_section and stripped == "地址：":
            _replace_paragraph_text(p, f"地址：{values.get('party_b_address', '')}")
            continue

        # Contract total amount
        if "合同总价为人民币：" in t:
            upper = values.get("contract_amount_upper", "")
            lower = values.get("contract_amount_lower", "")
            new_text = re.sub(
                r"人民币：\s*元整（￥\s*）",
                f"人民币：{upper}（￥{lower}）",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Tax rate
        if "合同总价为含税价（税率：" in t:
            rate = values.get("tax_rate", "")
            new_text = re.sub(
                r"税率：\s*%",
                f"税率：{rate}%",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Delivery time - after signing
        if "本合同签订后" in t and "天之内" in t:
            days = values.get("delivery_days_after_signing", "")
            new_text = re.sub(r"本合同签订后\s*天之内", f"本合同签订后{days}天之内", t)
            _replace_paragraph_text(p, new_text)
            continue

        # Delivery time - after prepayment
        if "收到预付款后" in t and "天之内" in t:
            days = values.get("delivery_days_after_prepayment", "")
            new_text = re.sub(r"收到预付款后\s*天之内", f"收到预付款后{days}天之内", t)
            _replace_paragraph_text(p, new_text)
            continue

        # Delivery location
        if "卖方将合同设备送到买方指定地点即" in t:
            location = values.get("delivery_location", "")
            new_text = re.sub(r"卖方将合同设备送到买方指定地点即\s*。", f"卖方将合同设备送到买方指定地点即{location}。", t)
            _replace_paragraph_text(p, new_text)
            continue

        # Prepayment
        if "合同签订后" in t and "的预付款" in t:
            days = values.get("prepayment_days", "")
            ratio = values.get("prepayment_ratio", "")
            upper = values.get("prepayment_amount_upper", "")
            lower = values.get("prepayment_amount_lower", "")
            new_text = re.sub(
                r"合同签订后\s*天内，买方支付合同总价\s*%即人民币\s*元整（￥\s*）的预付款",
                f"合同签订后{days}天内，买方支付合同总价{ratio}%即人民币{upper}（￥{lower}）的预付款",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Acceptance payment
        if "设备交付买方并安装调试" in t and "经买方终验收合格后" in t:
            days = values.get("acceptance_days", "")
            ratio = values.get("acceptance_ratio", "")
            upper = values.get("acceptance_amount_upper", "")
            lower = values.get("acceptance_amount_lower", "")
            new_text = re.sub(
                r"设备交付买方并安装调试，达到正常使用要求，经买方终验收合格后\s*天内，买方支付合同总价\s*%即人民币\s*元整（￥\s*）。",
                f"设备交付买方并安装调试，达到正常使用要求，经买方终验收合格后{days}天内，买方支付合同总价{ratio}%即人民币{upper}（￥{lower}）。",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Warranty payment
        if "剩余合同总价5%" in t:
            upper = values.get("warranty_amount_upper", "")
            lower = values.get("warranty_amount_lower", "")
            days = values.get("warranty_payment_days", "")
            new_text = re.sub(
                r"剩余合同总价5%即人民币\s*元整（￥\s*）作为质保金",
                f"剩余合同总价5%即人民币{upper}（￥{lower}）作为质保金",
                t,
            )
            new_text = re.sub(
                r"质保期满后\s*日之内",
                f"质保期满后{days}日之内",
                new_text,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Invoice type
        if "卖方收取合同价款之前应向买方提供" in t:
            invoice_type = values.get("invoice_type", "")
            if invoice_type == "invoice":
                _replace_paragraph_text(p, "卖方收取合同价款之前应向买方提供√合法等额完税发票 /     合法等额收据")
            elif invoice_type == "receipt":
                _replace_paragraph_text(p, "卖方收取合同价款之前应向买方提供    合法等额完税发票 / √合法等额收据")
            else:
                _replace_paragraph_text(p, t)
            continue

        # Warranty months
        if "设备的质量保证期为设备整机调试、试运行验收合格之日起" in t:
            months = values.get("warranty_months", "")
            new_text = re.sub(
                r"之日起\s*个月",
                f"之日起{months}个月",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Signing section: authorized representative
        if in_signing_section and stripped.startswith("授权代表："):
            if current_signing_party == "A":
                value = values.get("party_a_authorized_representative", "")
            else:
                value = values.get("party_b_authorized_representative", "")
            _replace_paragraph_text(p, f"授权代表：{value}")
            continue

        # Signing section: dates
        if in_signing_section and stripped == "日期：":
            if not party_a_date_filled:
                value = values.get("party_a_date", "")
                party_a_date_filled = True
            else:
                value = values.get("party_b_date", "")
            _replace_paragraph_text(p, f"日期：{value}")
            continue

        # Signing section: Party B legal representative
        if in_signing_section and stripped == "法定代表人：" and "申振威" not in t:
            _replace_paragraph_text(p, f"法定代表人：{values.get('party_b_legal_person', '')}")
            continue

    # Fill equipment table
    _fill_equipment_table(doc, equipment_items or [])

    doc.save(str(output_path))


def generate_filled_docx(
    template_path: Path,
    field_values: dict[str, str],
    equipment_items: list[dict[str, str]] | None = None,
) -> Path:
    """Generate a filled DOCX in a temporary file and return its path."""
    suffix = f"_{uuid.uuid4().hex}.docx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        output_path = Path(tmp.name)
    fill_contract(template_path, field_values, output_path, equipment_items=equipment_items)
    return output_path


def serialize_field_values(field_values: dict[str, str]) -> str:
    """Serialize field values to JSON for DB storage."""
    return json.dumps(field_values, ensure_ascii=False)


def deserialize_field_values(raw: str) -> dict[str, str]:
    """Deserialize field values from JSON DB storage."""
    return json.loads(raw)


def serialize_equipment_items(equipment_items: list[dict[str, str]] | None) -> str | None:
    """Serialize equipment items to JSON for DB storage."""
    if equipment_items is None:
        return None
    return json.dumps(equipment_items, ensure_ascii=False)


def deserialize_equipment_items(raw: str | None) -> list[dict[str, str]] | None:
    """Deserialize equipment items from JSON DB storage."""
    if raw is None:
        return None
    return json.loads(raw)
