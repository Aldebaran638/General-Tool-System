"""
Contract Filler DOCX Logic

Hardcoded filling rules for the construction contract template.
"""

import json
import re
import tempfile
import uuid
from pathlib import Path

from docx import Document

from .schemas import ContractField, PreviewSegment


# =============================================================================
# Hardcoded field definitions for the construction contract template
# =============================================================================

CONTRACT_FIELDS: list[ContractField] = [
    ContractField(key="party_b_name", label="承包方（乙方）名称"),
    ContractField(key="party_b_legal_person", label="乙方法定代表人"),
    ContractField(key="party_b_address", label="乙方地址"),
    ContractField(key="project_content", label="工程内容", type="textarea"),
    ContractField(key="total_duration_days", label="总工期（日历天）"),
    ContractField(key="contract_amount_upper", label="合同总价款大写"),
    ContractField(key="contract_amount_lower", label="合同总价款小写"),
    ContractField(key="prepayment_ratio", label="预付款比例（%）"),
    ContractField(key="prepayment_amount", label="预付款金额"),
    ContractField(key="acceptance_ratio", label="验收款比例（%）"),
    ContractField(key="acceptance_amount", label="验收款金额"),
    ContractField(key="warranty_ratio", label="质保金比例（%）"),
    ContractField(key="warranty_amount", label="质保金金额"),
    ContractField(key="party_a_contact_name", label="甲方现场负责人姓名"),
    ContractField(key="party_a_contact_phone", label="甲方现场负责人电话"),
    ContractField(key="party_b_contact_name", label="乙方现场负责人姓名"),
    ContractField(key="party_b_contact_phone", label="乙方现场负责人电话"),
    ContractField(key="attachment", label="附件"),
    ContractField(key="party_a_authorized_representative", label="甲方授权代表"),
    ContractField(key="party_a_date", label="甲方日期", type="date"),
    ContractField(key="party_b_authorized_representative", label="乙方授权代表"),
    ContractField(key="party_b_date", label="乙方日期", type="date"),
]

FIELD_KEYS = {f.key for f in CONTRACT_FIELDS}


# =============================================================================
# Preview extraction
# =============================================================================

# Mapping of regex patterns to field keys for inline marker insertion
_PREVIEW_MARKERS: list[tuple[str, str]] = [
    (r"(承包方（以下称乙方）：)\s*", r"\1{{party_b_name}}"),
    (r"^(法定代表人：)\s*$", r"\1{{party_b_legal_person}}"),
    (r"^(地址：)\s*$", r"\1{{party_b_address}}"),
    (r"(工程内容：)\s*", r"\1{{project_content}}"),
    (r"(自开工之日起)\s*(日历天)", r"\1{{total_duration_days}}\2"),
    (r"(人民币)\s*元整（小写￥\s*）", r"\1{{contract_amount_upper}}（小写￥{{contract_amount_lower}}）"),
    (r"(支付总工程款)\s*%即人民币\s*元整（￥\s*）的预付款", r"\1{{prepayment_ratio}}%即人民币{{prepayment_amount}}（￥{{prepayment_amount}}）的预付款"),
    (r"(支付总工程款)\s*%即人民币\s*元整（￥\s*）给乙方", r"\1{{acceptance_ratio}}%即人民币{{acceptance_amount}}（￥{{acceptance_amount}}）给乙方"),
    (r"(总工程款)\s*%即人民币\s*元整（￥\s*）给乙方", r"\1{{warranty_ratio}}%即人民币{{warranty_amount}}（￥{{warranty_amount}}）给乙方"),
    (r"(甲方派驻)\s*(为现场负责人，其联系电话)\s*", r"\1{{party_a_contact_name}}\2{{party_a_contact_phone}}"),
    (r"(乙方派驻)\s*(为工程现场负责人，其联系电话)\s*", r"\1{{party_b_contact_name}}\2{{party_b_contact_phone}}"),
    (r"(附件：)\s*", r"\1{{attachment}}"),
    (r"^(授权代表；|授权代表：)\s*$", r"\1{{party_a_authorized_representative}}"),
    (r"^(日期：)\s*$", r"\1{{party_a_date}}"),
    (r"^(承包方：)\s*", r"\1{{party_b_name}}"),
]


def _insert_preview_markers(text: str, *, in_signing_section: bool = False) -> str:
    """Replace blank placeholders in contract text with {{field_key}} markers."""
    for pattern, replacement in _PREVIEW_MARKERS:
        if in_signing_section and "乙方派驻" in pattern:
            continue
        if not in_signing_section and "承包方：" in replacement:
            continue
        new_text = re.sub(pattern, replacement, text, count=1)
        if new_text != text:
            return new_text
    return text


def extract_preview_with_markers(template_path: Path) -> list[PreviewSegment]:
    """Extract preview segments with {{field_key}} markers at blank positions."""
    doc = Document(str(template_path))
    segments: list[PreviewSegment] = []

    in_signing_section = False
    current_signing_party: str | None = None
    party_a_date_marker_done = False

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            continue

        if "发包方：" in text and "有利华建材" in text:
            in_signing_section = True
            current_signing_party = "A"
            segments.append(
                PreviewSegment(
                    type="paragraph",
                    text=text,
                    style=para.style.name if para.style else "Normal",
                )
            )
            continue

        # Signing section special handling for second date/authorized rep
        stripped = text.strip()
        if in_signing_section:
            if stripped.startswith("承包方："):
                current_signing_party = "B"
                text = re.sub(r"^(承包方：)\s*", r"\1{{party_b_name}}", text, count=1)
            elif stripped.startswith("授权代表"):
                if current_signing_party == "A":
                    text = re.sub(r"^(授权代表；|授权代表：)\s*", r"\1{{party_a_authorized_representative}}", text, count=1)
                else:
                    text = re.sub(r"^(授权代表；|授权代表：)\s*", r"\1{{party_b_authorized_representative}}", text, count=1)
            elif stripped == "日期：":
                if not party_a_date_marker_done:
                    text = re.sub(r"^(日期：)\s*", r"\1{{party_a_date}}", text, count=1)
                    party_a_date_marker_done = True
                else:
                    text = re.sub(r"^(日期：)\s*", r"\1{{party_b_date}}", text, count=1)
            elif stripped == "法定代表人：" and "申振威" not in text:
                text = re.sub(r"^(法定代表人：)\s*", r"\1{{party_b_legal_person}}", text, count=1)
            else:
                text = _insert_preview_markers(text, in_signing_section=True)
        else:
            text = _insert_preview_markers(text, in_signing_section=False)

        segments.append(
            PreviewSegment(
                type="paragraph",
                text=text,
                style=para.style.name if para.style else "Normal",
            )
        )

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        if rows:
            segments.append(PreviewSegment(type="table", rows=rows))

    return segments


def extract_preview(template_path: Path) -> list[PreviewSegment]:
    """Extract structured preview segments from the DOCX template."""
    return extract_preview_with_markers(template_path)


# =============================================================================
# Filling helpers
# =============================================================================

def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace all text in a paragraph while keeping the first run."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


# =============================================================================
# Main fill function
# =============================================================================

def fill_contract(template_path: Path, field_values: dict[str, str], output_path: Path) -> None:
    """Fill the construction contract template and save to output_path."""
    doc = Document(str(template_path))
    values = {k: v for k, v in field_values.items() if k in FIELD_KEYS}

    # State machine for signing section
    in_signing_section = False
    current_signing_party: str | None = None
    party_a_date_filled = False

    for p in doc.paragraphs:
        t = p.text
        stripped = t.strip()

        # Detect signing section start
        if "发包方：" in t and "有利华建材" in t:
            in_signing_section = True
            current_signing_party = "A"
            continue

        if in_signing_section and stripped == "承包方：":
            current_signing_party = "B"
            _replace_paragraph_text(
                p,
                f"承包方：{values.get('party_b_name', '')}",
            )
            continue

        # Party B name (body section)
        if not in_signing_section and "承包方（以下称乙方）：" in t:
            _replace_paragraph_text(
                p,
                f"承包方（以下称乙方）：{values.get('party_b_name', '')}",
            )
            continue

        # Party B legal representative (body section)
        if not in_signing_section and stripped == "法定代表人：" and "申振威" not in t:
            _replace_paragraph_text(p, f"法定代表人：{values.get('party_b_legal_person', '')}")
            continue

        # Party B address (body section)
        if not in_signing_section and stripped == "地址：":
            _replace_paragraph_text(p, f"地址：{values.get('party_b_address', '')}")
            continue

        # Project content
        if stripped == "工程内容：":
            _replace_paragraph_text(
                p,
                f"工程内容：{values.get('project_content', '')}",
            )
            continue

        # Total duration
        if "总工期为乙方自开工之日起" in t and "日历天" in t:
            days = values.get("total_duration_days", "")
            _replace_paragraph_text(p, t.replace("日历天", f"{days}日历天", 1))
            continue

        # Contract total amount
        if "本合同工程包干总价款为人民币" in t:
            upper = re.escape(values.get("contract_amount_upper", ""))
            lower = re.escape(values.get("contract_amount_lower", ""))
            new_text = re.sub(
                r"人民币\s*元整（小写￥\s*）",
                f"人民币{upper}（小写￥{lower}）",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Prepayment
        if "本合同签署生效后" in t and "的预付款给乙方" in t:
            ratio = re.escape(values.get("prepayment_ratio", ""))
            amount = re.escape(values.get("prepayment_amount", ""))
            new_text = re.sub(
                r"支付总工程款\s*%即人民币\s*元整（￥\s*）的预付款",
                f"支付总工程款{ratio}%即人民币{amount}（￥{amount}）的预付款",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Acceptance payment
        if "竣工验收合格后" in t and "给乙方。" in t and "质保金" not in t:
            ratio = re.escape(values.get("acceptance_ratio", ""))
            amount = re.escape(values.get("acceptance_amount", ""))
            new_text = re.sub(
                r"支付总工程款\s*%即人民币\s*元整（￥\s*）给乙方",
                f"支付总工程款{ratio}%即人民币{amount}（￥{amount}）给乙方",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Warranty payment
        if "保修期内无质量问题" in t:
            ratio = re.escape(values.get("warranty_ratio", ""))
            amount = re.escape(values.get("warranty_amount", ""))
            new_text = re.sub(
                r"总工程款\s*%即人民币\s*元整（￥\s*）给乙方",
                f"总工程款{ratio}%即人民币{amount}（￥{amount}）给乙方",
                t,
            )
            _replace_paragraph_text(p, new_text)
            continue

        # Party A contact
        if "甲方派驻" in t and "为现场负责人" in t:
            name = values.get("party_a_contact_name", "")
            phone = values.get("party_a_contact_phone", "")
            new_text = t.replace("甲方派驻", f"甲方派驻{name}", 1)
            new_text = new_text.replace("联系电话", f"联系电话{phone}", 1)
            _replace_paragraph_text(p, new_text)
            continue

        # Party B contact
        if "乙方派驻" in t and "为工程现场负责人" in t:
            name = values.get("party_b_contact_name", "")
            phone = values.get("party_b_contact_phone", "")
            new_text = t.replace("乙方派驻", f"乙方派驻{name}", 1)
            new_text = new_text.replace("联系电话", f"联系电话{phone}", 1)
            _replace_paragraph_text(p, new_text)
            continue

        # Tax selection line - keep selected
        if "含税价" in t and "税率" in t:
            _replace_paragraph_text(p, "√ 含税价（税率：13%）")
            continue

        # Attachment
        if stripped.startswith("附件："):
            _replace_paragraph_text(p, f"附件：{values.get('attachment', '')}")
            continue

        # Signing section: authorized representative
        if in_signing_section and (stripped == "授权代表；" or stripped == "授权代表："):
            if current_signing_party == "A":
                value = values.get("party_a_authorized_representative", "")
            else:
                value = values.get("party_b_authorized_representative", "")
            _replace_paragraph_text(p, f"授权代表：{value}")
            continue

        # Signing section: dates
        if in_signing_section and stripped == "日期：":
            if not party_a_date_filled:
                value = values.get("party_a_date", "")
                party_a_date_filled = True
            else:
                value = values.get("party_b_date", "")
            _replace_paragraph_text(p, f"日期：{value}")
            continue

        # Signing section: Party B legal representative
        if in_signing_section and stripped == "法定代表人：" and "申振威" not in t:
            _replace_paragraph_text(
                p,
                f"法定代表人：{values.get('party_b_legal_person', '')}",
            )
            continue

    doc.save(str(output_path))


def generate_filled_docx(template_path: Path, field_values: dict[str, str]) -> Path:
    """Generate a filled DOCX in a temporary file and return its path."""
    suffix = f"_{uuid.uuid4().hex}.docx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        output_path = Path(tmp.name)
    fill_contract(template_path, field_values, output_path)
    return output_path


def serialize_field_values(field_values: dict[str, str]) -> str:
    """Serialize field values to JSON for DB storage."""
    return json.dumps(field_values, ensure_ascii=False)


def deserialize_field_values(raw: str) -> dict[str, str]:
    """Deserialize field values from JSON DB storage."""
    return json.loads(raw)
