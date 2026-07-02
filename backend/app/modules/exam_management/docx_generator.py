"""
Exam Management — Docx Generator Service

Generates .docx exam papers from exam questions and options.
"""

from __future__ import annotations

import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlmodel import Session, select

from app.core.config import settings
from app.core.storage import resolve_papers_dir
from app.modules.exam_management.models import (
    Exam,
    Question,
    QuestionOption,
)


_TYPE_LABELS: dict[str, str] = {
    "SINGLE_CHOICE": "单选题",
    "MULTIPLE_CHOICE": "多选题",
    "TRUE_FALSE": "判断题",
}


def generate_exam_paper_docx(exam_id: uuid.UUID, session: Session) -> str:
    """Generate a .docx exam paper and return the file path.

    Args:
        exam_id: The exam UUID.
        session: SQLModel database session.

    Returns:
        Relative path to the generated .docx file (relative to upload root).

    Raises:
        ValueError: If the exam does not exist.
    """
    exam = session.get(Exam, exam_id)
    if not exam:
        raise ValueError("考试不存在")

    # Query questions ordered by sort_no
    questions = session.exec(
        select(Question)
        .where(Question.exam_id == exam_id)
        .order_by(Question.sort_no)
    ).all()

    # Batch-load all options to avoid N+1 queries
    question_ids = [q.id for q in questions]
    all_options = session.exec(
        select(QuestionOption)
        .where(QuestionOption.question_id.in_(question_ids))
        .order_by(QuestionOption.sort_no)
    ).all()
    options_by_question: dict[uuid.UUID, list[QuestionOption]] = {}
    for opt in all_options:
        options_by_question.setdefault(opt.question_id, []).append(opt)

    # Build document
    doc = Document()

    # Title
    title = doc.add_heading(exam.name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Exam meta
    doc.add_paragraph(f"考试时长: {exam.duration_minutes} 分钟")
    doc.add_paragraph(f"及格分数: {exam.pass_score} 分")
    doc.add_paragraph("")  # blank line

    # Questions
    for i, q in enumerate(questions, 1):
        type_label = _TYPE_LABELS.get(q.question_type, q.question_type)
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. [{type_label}] {q.stem}  ({q.score}分)")
        run.bold = True

        options = options_by_question.get(q.id, [])
        for opt in options:
            prefix = "✓ " if opt.is_correct else "   "
            doc.add_paragraph(f"{prefix}{opt.option_key}. {opt.option_text}")

        if q.analysis:
            p_analysis = doc.add_paragraph()
            run_a = p_analysis.add_run(f"解析: {q.analysis}")
            run_a.italic = True

        doc.add_paragraph("")  # separator

    # Save file
    papers_dir = resolve_papers_dir()
    papers_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{exam_id}_{uuid.uuid4().hex[:8]}.docx"
    filepath = papers_dir / filename
    doc.save(str(filepath))

    # Return path relative to upload root for static file serving
    return f"papers/{filename}"
